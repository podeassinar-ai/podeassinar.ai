#!/usr/bin/env python3
"""
Document Text Extraction Script

This script extracts text and tables from documents using deterministic libraries:
- pdfplumber for native PDFs
- pytesseract (Tesseract OCR) for scanned PDFs and images
- python-docx for Word documents

Usage:
    python extractor.py --input <base64_content> --mime-type <mime_type>

Output:
    JSON object with extracted text, tables, and metadata
"""

import argparse
import base64
import io
import json
import sys
from datetime import datetime
from typing import Any

# Attempt imports with graceful fallbacks
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from PIL import Image
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_from_pdf(content: bytes) -> dict[str, Any]:
    """Extract text and tables from PDF files."""
    if not HAS_PDFPLUMBER:
        return {
            "success": False,
            "error": "pdfplumber not installed. Run: pip install pdfplumber",
            "text": "",
            "tables": [],
            "metadata": {}
        }
    
    text_parts = []
    tables = []
    warnings = []
    used_ocr = False
    page_count = 0
    
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            
            for i, page in enumerate(pdf.pages):
                # Try native text extraction first
                page_text = page.extract_text() or ""
                
                # If no text found, try OCR
                if not page_text.strip() and HAS_TESSERACT:
                    used_ocr = True
                    try:
                        # Convert page to image for OCR
                        img = page.to_image(resolution=300)
                        pil_image = img.original
                        page_text = pytesseract.image_to_string(pil_image, lang='por')
                    except Exception as ocr_err:
                        warnings.append(f"OCR failed on page {i+1}: {str(ocr_err)}")
                
                if page_text.strip():
                    text_parts.append(f"## Página {i+1}\n\n{page_text.strip()}")
                
                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        tables.append({
                            "title": f"Tabela da página {i+1}",
                            "data": table
                        })
        
        return {
            "success": True,
            "text": "\n\n".join(text_parts),
            "tables": tables,
            "metadata": {
                "pageCount": page_count,
                "usedOcr": used_ocr,
                "extractedAt": datetime.now().isoformat(),
                "warnings": warnings if warnings else None
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"PDF extraction failed: {str(e)}",
            "text": "",
            "tables": [],
            "metadata": {"usedOcr": used_ocr}
        }


def extract_from_image(content: bytes) -> dict[str, Any]:
    """Extract text from images using OCR."""
    if not HAS_TESSERACT:
        return {
            "success": False,
            "error": "pytesseract not installed. Run: pip install pytesseract Pillow",
            "text": "",
            "tables": [],
            "metadata": {}
        }
    
    try:
        image = Image.open(io.BytesIO(content))
        
        # Get OCR data with confidence
        ocr_data = pytesseract.image_to_data(image, lang='por', output_type=pytesseract.Output.DICT)
        
        # Calculate average confidence
        confidences = [c for c in ocr_data['conf'] if c > 0]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        # Extract text
        text = pytesseract.image_to_string(image, lang='por')
        
        return {
            "success": True,
            "text": text.strip(),
            "tables": [],
            "metadata": {
                "usedOcr": True,
                "ocrConfidence": round(avg_confidence / 100, 2),
                "extractedAt": datetime.now().isoformat()
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Image OCR failed: {str(e)}",
            "text": "",
            "tables": [],
            "metadata": {"usedOcr": True}
        }


def extract_from_docx(content: bytes) -> dict[str, Any]:
    """Extract text from Word documents."""
    if not HAS_DOCX:
        return {
            "success": False,
            "error": "python-docx not installed. Run: pip install python-docx",
            "text": "",
            "tables": [],
            "metadata": {}
        }
    
    try:
        doc = DocxDocument(io.BytesIO(content))
        
        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append({
                    "title": f"Tabela {i+1}",
                    "data": table_data
                })
        
        return {
            "success": True,
            "text": "\n\n".join(paragraphs),
            "tables": tables,
            "metadata": {
                "usedOcr": False,
                "extractedAt": datetime.now().isoformat()
            }
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"DOCX extraction failed: {str(e)}",
            "text": "",
            "tables": [],
            "metadata": {"usedOcr": False}
        }


def extract_plain_text(content: bytes) -> dict[str, Any]:
    """Extract content from plain text files."""
    try:
        text = content.decode('utf-8')
        return {
            "success": True,
            "text": text,
            "tables": [],
            "metadata": {
                "usedOcr": False,
                "extractedAt": datetime.now().isoformat()
            }
        }
    except UnicodeDecodeError:
        try:
            text = content.decode('latin-1')
            return {
                "success": True,
                "text": text,
                "tables": [],
                "metadata": {
                    "usedOcr": False,
                    "extractedAt": datetime.now().isoformat(),
                    "warnings": ["Decoded using latin-1 fallback encoding"]
                }
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Text decoding failed: {str(e)}",
                "text": "",
                "tables": [],
                "metadata": {"usedOcr": False}
            }


def main():
    parser = argparse.ArgumentParser(description='Extract text from documents')
    parser.add_argument('--input', required=True, help='Base64 encoded file content')
    parser.add_argument('--mime-type', required=True, help='MIME type of the file')
    args = parser.parse_args()
    
    try:
        content = base64.b64decode(args.input)
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"Failed to decode base64 input: {str(e)}",
            "text": "",
            "tables": [],
            "metadata": {}
        }))
        sys.exit(1)
    
    mime_type = args.mime_type.lower()
    
    # Route to appropriate extractor
    if mime_type == 'application/pdf':
        result = extract_from_pdf(content)
    elif mime_type in ('image/png', 'image/jpeg', 'image/jpg', 'image/webp', 'image/tiff'):
        result = extract_from_image(content)
    elif mime_type in ('application/vnd.openxmlformats-officedocument.wordprocessingml.document',):
        result = extract_from_docx(content)
    elif mime_type in ('text/plain', 'text/markdown'):
        result = extract_plain_text(content)
    else:
        # Fallback: try plain text
        result = extract_plain_text(content)
        if not result["success"]:
            result = {
                "success": False,
                "error": f"Unsupported MIME type: {mime_type}",
                "text": "",
                "tables": [],
                "metadata": {"mimeType": mime_type}
            }
    
    # Add mime type to metadata
    result["metadata"]["mimeType"] = mime_type
    
    print(json.dumps(result, ensure_ascii=False))


if __name__ == '__main__':
    main()
